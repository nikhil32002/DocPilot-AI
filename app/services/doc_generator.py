from pathlib import Path

from docx import Document


class DocumentGenerator:

    OUTPUT_DIR = Path("app/outputs")

    def generate(
            self,
            title: str,
            sections: dict[str, str]
    ):

        self.OUTPUT_DIR.mkdir(exist_ok=True)

        document = Document()

        document.add_heading(title, level=1)

        for heading, content in sections.items():

            document.add_heading(heading, level=2)

            document.add_paragraph(content)

        filename = (
            title
            .replace(" ", "_")
            .replace("/", "")
        )

        output = self.OUTPUT_DIR / f"{filename}.docx"

        document.save(output)

        return str(output)


document_generator = DocumentGenerator()